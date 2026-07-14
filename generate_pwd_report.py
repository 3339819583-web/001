"""生成密码修改漏洞修复报告"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def shading(cell, color):
    s = cell._element.get_or_add_tcPr()
    e = s.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color})
    s.append(e)


def code(doc, text, color=RGBColor(0x8E, 0x44, 0xAD)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = color


def make_table(doc, headers, data):
    t = doc.add_table(rows=len(data)+1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        shading(c, "2C3E50")
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
            if "✅" in val:
                shading(t.rows[i].cells[j], "E8F8F5")
            elif "❌" in val:
                shading(t.rows[i].cells[j], "FDEDEC")
    return t


doc = Document()
t = doc.add_heading("密码修改漏洞修复报告", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("项目：Flask 用户信息管理平台\n").font.size = Pt(11)
p.add_run("修复日期：2026-07-13\n").font.size = Pt(11)
r = p.add_run("修复状态：已全部修复")
r.bold = True
r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
doc.add_paragraph()

doc.add_heading("一、漏洞概述", level=1)
doc.add_paragraph("密码修改功能存在三个严重安全漏洞，攻击者可以任意修改其他用户的密码。")

doc.add_heading("二、漏洞详情", level=1)

doc.add_heading("漏洞 1：越权修改密码", level=2)
doc.add_paragraph("风险等级：🔴 严重", style="List Bullet")
p = doc.add_paragraph()
r = p.add_run("问题代码：")
r.bold = True
r.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
code(doc, 'target_username = request.form.get("username", "")')
code(doc, '# 直接用表单传来的 username，未与 session 比对')
code(doc, 'USERS[target_username]["password"] = new_password')
p = doc.add_paragraph()
r = p.add_run("修复代码：")
r.bold = True
r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
code(doc, '# 只能修改自己的密码')
code(doc, 'if target_username != current_username:')
code(doc, '    return redirect(...)')
doc.add_paragraph("攻击场景：alice 登录后把表单 username 改成 admin，即可修改 admin 密码")

doc.add_heading("漏洞 2：无 CSRF 保护", level=2)
doc.add_paragraph("风险等级：🔴 严重", style="List Bullet")
p = doc.add_paragraph()
r = p.add_run("修复代码：")
r.bold = True
r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
code(doc, 'token = request.form.get("_csrf_token", "")')
code(doc, 'if not token or token != session.get("_csrf_token"):')
code(doc, '    abort(403)')

doc.add_heading("漏洞 3：无需原密码", level=2)
doc.add_paragraph("风险等级：🟠 高危", style="List Bullet")
doc.add_paragraph("攻击场景：攻击者只要拿到用户的一次登录 session，即可永久修改密码", style="List Bullet")

doc.add_heading("三、修复前后对比", level=1)
make_table(doc, ["安全措施", "修复前", "修复后"], [
    ["身份校验", "❌ 不验证 session", "✅ 只能改自己的"],
    ["CSRF 保护", "❌ 无", "✅ CSRF Token 校验"],
    ["原密码验证", "❌ 不需要", "✅ 需要原密码"],
])

doc.add_paragraph()
doc.add_heading("四、修复验证", level=1)
make_table(doc, ["测试项", "结果"], [
    ["无 CSRF Token 请求", "✅ 403 拦截"],
    ["alice 改 admin 密码", "✅ 拦截：只能改自己"],
    ["原密码错误时改密", "✅ 拦截：原密码错误"],
    ["正常修改密码", "✅ 成功"],
])

doc.add_paragraph()
doc.add_heading("五、版本历史", level=1)
make_table(doc, ["版本", "日期", "变更", "状态"], [
    ["v1.0", "07-07", "初始版本", "❌"],
    ["v1.1", "07-07", "密码/CSRF加固", "⚠️"],
    ["v1.2", "07-08", "修复SQL注入", "⚠️"],
    ["v1.3", "07-09", "修复文件上传", "⚠️"],
    ["v1.4", "07-09", "修复越权", "⚠️"],
    ["v1.5", "07-13", "修复LFI", "⚠️"],
    ["v1.6", "07-13", "修复密码修改漏洞", "✅"],
])

path = "密码修改漏洞修复报告.docx"
doc.save(path)
print(f"✅ {path}")
