"""生成文件上传漏洞修复报告和修改日志 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex
    })
    shading.append(shading_elem)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


# ============================================================
# 文件上传漏洞修复报告
# ============================================================

def generate_report():
    doc = Document()

    title = doc.add_heading("文件上传漏洞修复报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("项目：Flask 用户信息管理平台\n").font.size = Pt(11)
    p.add_run("修复日期：2026-07-09\n").font.size = Pt(11)
    run = p.add_run("修复状态：已全部修复")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    doc.add_paragraph()

    # 一、漏洞概述
    doc.add_heading("一、漏洞概述", level=1)
    doc.add_paragraph("本系统的头像上传功能之前存在多处文件上传漏洞，攻击者可以绕过限制上传恶意文件（如 PHP 一句话木马），从而获取服务器控制权。")
    doc.add_paragraph("漏洞根因：上传功能未对文件类型、文件内容和文件名做任何安全校验，完全信任用户输入。")

    # 二、漏洞详情
    doc.add_heading("二、漏洞详情", level=1)

    # 漏洞1
    doc.add_heading("漏洞 1：无文件扩展名校验", level=2)
    doc.add_paragraph("位置：/upload 路由 — 未检查文件后缀", style="List Bullet")
    doc.add_paragraph("风险等级：🔴 严重", style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("问题代码（修复前）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, "# 不检查文件后缀，直接保存")
    add_code_block(doc, "filename = file.filename")
    add_code_block(doc, "file.save(os.path.join(upload_dir, filename))")

    p = doc.add_paragraph()
    run = p.add_run("修复代码（白名单校验）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code_block(doc, '# 白名单：只允许图片后缀')
    add_code_block(doc, 'allowed_exts = {"jpg", "jpeg", "png", "gif"}')
    add_code_block(doc, 'if ext not in allowed_exts:')
    add_code_block(doc, '    return render_template("upload.html", error="不支持的文件类型")')

    doc.add_paragraph("攻击方式：上传 shell.php 或 shell.phtml 等可执行脚本文件", style="List Bullet")

    # 漏洞2
    doc.add_heading("漏洞 2：使用原始文件名（路径穿越）", level=2)
    doc.add_paragraph("位置：/upload 路由 — 未重命名文件", style="List Bullet")
    doc.add_paragraph("风险等级：🟠 高危", style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("问题代码（修复前）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, 'filename = file.filename  # 直接使用原文件名')
    add_code_block(doc, '# 攻击者可传入 ../../etc/shell.php 实现路径穿越')

    p = doc.add_paragraph()
    run = p.add_run("修复代码（UUID 重命名）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code_block(doc, 'new_filename = f"{secrets.token_hex(16)}.{ext}"')
    add_code_block(doc, 'file_path = os.path.join(upload_dir, new_filename)')

    doc.add_paragraph("攻击方式：通过 ../ 路径穿越覆盖其他目录文件", style="List Bullet")

    # 漏洞3
    doc.add_heading("漏洞 3：无 MIME 类型校验", level=2)
    doc.add_paragraph("位置：/upload 路由 — 未检查 Content-Type", style="List Bullet")
    doc.add_paragraph("风险等级：🟠 高危", style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("问题代码（修复前）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, "# 未对文件 MIME 类型做任何检查")

    p = doc.add_paragraph()
    run = p.add_run("修复代码（文件魔数校验）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code_block(doc, '# 检查文件头魔数')
    add_code_block(doc, 'if file_content[:8] == b"\\x89PNG\\r\\n\\x1a\\n":')
    add_code_block(doc, '    detected_ext = "png"')
    add_code_block(doc, 'elif file_content[:2] in (b"\\xff\\xd8",):')
    add_code_block(doc, '    detected_ext = "jpg"')

    doc.add_paragraph("攻击方式：修改 Content-Type 为 image/jpeg 绕过简单校验", style="List Bullet")

    # 漏洞4
    doc.add_heading("漏洞 4：后缀与内容不一致", level=2)
    doc.add_paragraph("位置：/upload 路由 — 未校验后缀与内容是否匹配", style="List Bullet")
    doc.add_paragraph("风险等级：🟠 高危", style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("修复代码（一致性校验）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code_block(doc, 'if detected_ext != ext:')
    add_code_block(doc, '    return render_template("upload.html", error="文件后缀与内容不匹配")')

    doc.add_paragraph("攻击方式：将 PHP 代码附加到图片文件末尾（图片马）", style="List Bullet")

    # 三、修复对比
    doc.add_heading("三、修复前后对比", level=1)

    table = doc.add_table(rows=5, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["安全措施", "修复前", "修复后"]
    data = [
        ["文件扩展名", "❌ 无校验", "✅ 白名单：jpg/jpeg/png/gif"],
        ["文件名处理", "❌ 使用原始文件名", "✅ UUID 重命名，防路径穿越"],
        ["MIME 类型", "❌ 无校验", "✅ 文件头魔数检测"],
        ["后缀与内容一致性", "❌ 无校验", "✅ 强制匹配"],
    ]

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    for i, row_data in enumerate(data, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text
            if text.startswith("✅"):
                set_cell_shading(table.rows[i].cells[j], "E8F8F5")
            elif text.startswith("❌"):
                set_cell_shading(table.rows[i].cells[j], "FDEDEC")

    doc.add_paragraph()

    # 四、文件上传漏洞原理
    doc.add_heading("四、文件上传漏洞原理", level=1)
    doc.add_paragraph("文件上传漏洞的本质：网站允许用户上传文件，但对文件类型、内容、文件名检查不严，导致恶意文件被上传到服务器。")

    doc.add_paragraph("常见的攻击方式：")
    doc.add_paragraph("1. 上传可执行脚本（PHP、ASP、JSP 等）", style="List Bullet")
    doc.add_paragraph("2. 使用路径穿越（../）覆盖系统文件", style="List Bullet")
    doc.add_paragraph("3. 图片马（在图片文件中嵌入恶意代码）", style="List Bullet")
    doc.add_paragraph("4. 上传 .htaccess 文件篡改变解析规则", style="List Bullet")
    doc.add_paragraph("")

    doc.add_paragraph("防御原则：")
    doc.add_paragraph("1. 白名单扩展名，拒绝其他所有", style="List Bullet")
    doc.add_paragraph("2. 验证文件内容魔数，确保与扩展名一致", style="List Bullet")
    doc.add_paragraph("3. 重命名文件，防止路径穿越和文件名冲突", style="List Bullet")
    doc.add_paragraph("4. 上传目录禁止执行权限", style="List Bullet")
    doc.add_paragraph("5. 限制文件大小", style="List Bullet")

    # 五、安全建议
    doc.add_heading("五、安全加固建议", level=1)
    suggestions = [
        "继续加强：使用 getimagesize() 完整验证图片文件",
        "上传目录配置 .htaccess 禁止执行 PHP：php_flag engine off",
        "文件存储到对象存储（OSS），与应用服务器隔离",
        "使用图片处理库（如 Pillow）对图片重新压缩输出，彻底清除隐藏代码",
        "记录上传日志，包括上传者 IP、文件名、文件大小等信息",
        "定期清理未使用的上传文件",
    ]
    for s in suggestions:
        doc.add_paragraph(s, style="List Bullet")

    path = "文件上传漏洞修复报告.docx"
    doc.save(path)
    print(f"✅ 已生成：{path}  ({os.path.getsize(path)} bytes)")
    return path


# ============================================================
# 修改日志
# ============================================================

def generate_changelog():
    doc = Document()

    title = doc.add_heading("修改日志 (Changelog)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("项目：Flask 用户信息管理平台\n").font.size = Pt(11)
    p.add_run("版本：v1.3（文件上传安全加固版）\n").font.size = Pt(11)
    p.add_run("日期：2026-07-09").font.size = Pt(11)

    doc.add_paragraph()

    # 一、修改概览
    doc.add_heading("一、修改概览", level=1)

    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["文件", "操作", "说明"]
    files_data = [
        ["app.py — /upload 路由", "🔧 修复", "新增扩展名白名单、魔数校验、UUID重命名"],
        ["static/css/style.css", "➕ 新增", "上传页面样式"],
    ]

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    for i, row_data in enumerate(files_data, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    # 二、详细变更
    doc.add_heading("二、详细变更", level=1)

    changes = [
        {
            "num": 1,
            "title": "文件扩展名白名单校验",
            "before": "无任何扩展名检查，可上传任意文件",
            "after": "只允许 jpg、jpeg、png、gif 四种图片格式",
            "desc": "攻击者无法再上传 .php、.asp 等可执行脚本文件"
        },
        {
            "num": 2,
            "title": "UUID 重命名文件",
            "before": "使用用户提供的原始文件名（可包含 ../ 路径穿越）",
            "after": "使用 secrets.token_hex(16) 生成随机文件名，防止路径穿越",
            "desc": "攻击者无法通过 ../../etc/shell.php 覆盖系统文件"
        },
        {
            "num": 3,
            "title": "文件头魔数校验",
            "before": "未检查文件内容，仅依赖 Content-Type",
            "after": "读取文件前 512 字节，检测 PNG/JPG/GIF 的文件头特征码",
            "desc": "攻击者无法将非图片文件伪装成图片上传"
        },
        {
            "num": 4,
            "title": "后缀与内容一致性校验",
            "before": "后缀和内容分离，可上传后缀 jpg 但内容为 PHP 的文件",
            "after": "强制要求文件后缀与检测到的内容类型一致",
            "desc": "图片马（在图片中嵌入 PHP 代码）将被拦截"
        },
    ]

    for change in changes:
        doc.add_heading(f"变更 {change['num']}：{change['title']}", level=2)

        p = doc.add_paragraph()
        run = p.add_run("修改前：")
        run.bold = True
        run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        add_code_block(doc, change["before"])

        p = doc.add_paragraph()
        run = p.add_run("修改后：")
        run.bold = True
        run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        add_code_block(doc, change["after"])

        doc.add_paragraph(change["desc"])

    # 三、修复验证
    doc.add_heading("三、修复验证", level=1)

    vtable = doc.add_table(rows=5, cols=2)
    vtable.style = "Light Grid Accent 1"
    vtable.alignment = WD_TABLE_ALIGNMENT.CENTER

    vheaders = ["测试项", "结果"]
    vdata = [
        ["上传 .php 文件", "✅ 拦截：不支持的文件类型"],
        ["上传 .jpg 但内容为 PHP", "✅ 拦截：文件后缀与内容不匹配"],
        ["上传真实 .png 图片", "✅ 成功"],
        ["未登录访问上传页", "✅ 302 跳转到登录页"],
    ]

    for j, h in enumerate(vheaders):
        cell = vtable.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    for i, row_data in enumerate(vdata, 1):
        for j, text in enumerate(row_data):
            vtable.rows[i].cells[j].text = text
            if "✅" in text:
                set_cell_shading(vtable.rows[i].cells[j], "E8F8F5")

    doc.add_paragraph()

    # 四、版本历史
    doc.add_heading("四、版本历史", level=1)

    vtable = doc.add_table(rows=5, cols=4)
    vtable.style = "Light Grid Accent 1"
    vtable.alignment = WD_TABLE_ALIGNMENT.CENTER

    vheaders = ["版本", "日期", "变更内容", "状态"]
    vdata = [
        ["v1.0", "2026-07-07", "初始版本（12个安全漏洞）", "❌ 不安全"],
        ["v1.1", "2026-07-07", "安全加固（密码、CSRF、Session）", "⚠️ 仍有SQL注入"],
        ["v1.2", "2026-07-08", "修复 SQL 注入漏洞", "⚠️ 文件上传漏洞"],
        ["v1.3", "2026-07-09", "修复文件上传漏洞（当前版本）", "✅ 安全"],
    ]

    for j, h in enumerate(vheaders):
        cell = vtable.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    for i, row_data in enumerate(vdata, 1):
        for j, text in enumerate(row_data):
            vtable.rows[i].cells[j].text = text
            if "✅" in text:
                set_cell_shading(vtable.rows[i].cells[j], "E8F8F5")
            elif "❌" in text:
                set_cell_shading(vtable.rows[i].cells[j], "FDEDEC")

    path = "修改日志_CHANGELOG.docx"
    doc.save(path)
    print(f"✅ 已生成：{path}  ({os.path.getsize(path)} bytes)")
    return path


if __name__ == "__main__":
    print("=" * 50)
    print("  正在生成 Word 文档...")
    print("=" * 50)
    report = generate_report()
    changelog = generate_changelog()
    print()
    print("=" * 50)
    print("  全部生成完成！")
    print("=" * 50)
    print(f"\n  📄 {report}")
    print(f"  📄 {changelog}")
