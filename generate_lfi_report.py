"""生成文件包含漏洞修复报告"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex
    })
    shading.append(shading_elem)


def add_code(doc, text, color=RGBColor(0x8E, 0x44, 0xAD)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = color
    return p


def generate_report():
    doc = Document()
    t = doc.add_heading("文件包含漏洞修复报告", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("项目：Flask 用户信息管理平台\n").font.size = Pt(11)
    p.add_run("修复日期：2026-07-13\n").font.size = Pt(11)
    r = p.add_run("修复状态：已全部修复")
    r.bold = True
    r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    doc.add_paragraph()

    # 漏洞概述
    doc.add_heading("一、漏洞概述", level=1)
    doc.add_paragraph("本系统的动态页面加载功能存在本地文件包含（LFI）漏洞，攻击者可以利用路径穿越（../）读取服务器上的任意文件。")
    doc.add_paragraph("漏洞根因：os.path.join() 拼接用户输入时，未对 ../ 等路径穿越字符做任何过滤，导致攻击者可以跳出 pages/ 目录读取系统文件。")

    # 漏洞详情
    doc.add_heading("二、漏洞详情", level=1)

    doc.add_heading("漏洞：动态页面 LFI（本地文件包含）", level=2)
    doc.add_paragraph("风险等级：🔴 严重", style="List Bullet")
    doc.add_paragraph("位置：/page 路由 — name 参数", style="List Bullet")

    r = doc.add_paragraph()
    run = r.add_run("问题代码（修复前）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code(doc, 'name = request.args.get("name", "")')
    add_code(doc, 'filepath = os.path.join("pages", name)   # ← 直接拼接，不过滤')
    add_code(doc, 'open(filepath).read()')

    r = doc.add_paragraph()
    run = r.add_run("修复代码（白名单校验）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code(doc, '# 白名单：只允许预定义的页面')
    add_code(doc, 'allowed_pages = {"help", "about"}')
    add_code(doc, '# 验证：拒绝非字母数字和不在白名单的输入')
    add_code(doc, 'if not name.isalnum() or name not in allowed_pages:')
    add_code(doc, '    return render_template("...", page_content="页面不存在")')

    doc.add_paragraph("攻击场景：")
    doc.add_paragraph("访问 /page?name=../../../../etc/passwd → 读取 Linux 系统账号文件", style="List Bullet")
    doc.add_paragraph("访问 /page?name=../app.py → 读取 Flask 应用源代码", style="List Bullet")
    doc.add_paragraph("访问 /page?name=../data/users.db → 下载 SQLite 数据库文件", style="List Bullet")

    # 修复对比
    doc.add_heading("三、修复前后对比", level=1)
    tbl1 = doc.add_table(rows=4, cols=3)
    tbl1.style = "Light Grid Accent 1"
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["对比项", "修复前", "修复后"]):
        c = tbl1.rows[0].cells[j]
        c.text = h
        set_cell_shading(c, "2C3E50")
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
    for i, d in enumerate([
        ["路径校验", "❌ 无校验，直接拼接", "✅ 白名单 + isalnum()"],
        ["../ 穿越", "❌ 允许穿越到系统目录", "✅ 拒绝非字母数字"],
        ["读取 /etc/passwd", "✅ 可读取", "❌ 拦截"],
    ], 1):
        for j, txt in enumerate(d):
            c = tbl1.rows[i].cells[j]
            c.text = txt
            if "✅" in txt or "❌" in txt:
                set_cell_shading(c, "E8F8F5" if "✅" in txt else "FDEDEC")
    doc.add_paragraph()

    # 修复验证
    doc.add_heading("四、修复验证", level=1)
    tbl2 = doc.add_table(rows=5, cols=2)
    tbl2.style = "Light Grid Accent 1"
    for j, h in enumerate(["测试项", "结果"]):
        c = tbl2.rows[0].cells[j]
        c.text = h
        set_cell_shading(c, "2C3E50")
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
    for i, (test, result) in enumerate([
        ("正常访问 help", "✅ 成功加载"),
        ("../ 穿越读 /etc/passwd", "✅ 拦截：页面不存在"),
        ("../ 穿越读 app.py", "✅ 拦截：页面不存在"),
        ("非法字符注入", "✅ 拦截：页面不存在"),
    ], 1):
        tbl2.rows[i].cells[0].text = test
        tbl2.rows[i].cells[1].text = result
        if "✅" in result:
            set_cell_shading(tbl2.rows[i].cells[1], "E8F8F5")

    doc.add_paragraph()

    # LFI 原理
    doc.add_heading("五、LFI 漏洞原理", level=1)
    doc.add_paragraph("LFI（Local File Inclusion）本地文件包含漏洞的原理：")
    doc.add_paragraph("")
    doc.add_paragraph("当应用程序使用用户输入拼接文件路径时，如果未做安全校验，攻击者可以利用 ../ 目录回溯符号跳出限制目录，读取系统上的任意文件。")
    doc.add_paragraph("")
    doc.add_paragraph("路径解析示例：")
    add_code(doc, "用户输入：name=../../../../etc/passwd")
    add_code(doc, "拼接后：pages/../../../../etc/passwd")
    add_code(doc, "解析后：/etc/passwd                    ← 读到系统文件！")
    doc.add_paragraph("")
    doc.add_paragraph("防御原则：")
    doc.add_paragraph("1. 白名单校验：只允许预定义的页面名称", style="List Bullet")
    doc.add_paragraph("2. 过滤 ../：拒绝包含路径穿越字符的输入", style="List Bullet")
    doc.add_paragraph("3. 使用 os.path.realpath()：规范化路径后比较前缀", style="List Bullet")

    # 安全建议
    doc.add_heading("六、安全加固建议", level=1)
    for s in [
        "始终使用白名单控制可访问的文件",
        "对用户输入做严格校验：只允许字母数字",
        "使用 os.path.realpath() 规范化路径后校验前缀",
        "避免使用用户输入直接拼接文件路径",
        "部署 WAF 拦截路径穿越攻击（../）",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    path = "文件包含漏洞修复报告.docx"
    doc.save(path)
    print(f"✅ {path}  ({os.path.getsize(path)} bytes)")


if __name__ == "__main__":
    generate_report()
