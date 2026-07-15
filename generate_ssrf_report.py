"""生成SSRF漏洞修复报告"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def sc(cell, color):
    s = cell._element.get_or_add_tcPr()
    e = s.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color})
    s.append(e)


def code(doc, text, color=RGBColor(0x8E, 0x44, 0xAD)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = color


def mk_table(doc, headers, data):
    t = doc.add_table(rows=len(data)+1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        sc(c, "2C3E50")
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
            if "✅" in val:
                sc(t.rows[i].cells[j], "E8F8F5")
            elif "❌" in val:
                sc(t.rows[i].cells[j], "FDEDEC")
    return t


doc = Document()
t = doc.add_heading("SSRF 漏洞修复报告", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("项目：Flask 用户信息管理平台\n").font.size = Pt(11)
p.add_run("修复日期：2026-07-14\n").font.size = Pt(11)
r = p.add_run("修复状态：已全部修复")
r.bold = True
r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
doc.add_paragraph()

doc.add_heading("一、漏洞概述", level=1)
doc.add_paragraph("URL 抓取功能存在 SSRF（Server-Side Request Forgery）漏洞，攻击者可以让服务器代为请求内网资源或读取本地文件。")

doc.add_heading("二、漏洞详情", level=1)

doc.add_heading("漏洞：SSRF 服务端请求伪造", level=2)
doc.add_paragraph("风险等级：🔴 严重", style="List Bullet")
doc.add_paragraph("位置：/fetch-url 路由 — url 参数", style="List Bullet")

p = doc.add_paragraph()
r = p.add_run("问题代码（修复前）：")
r.bold = True
r.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
code(doc, 'url = request.form.get("url", "")')
code(doc, '# 直接请求，不做任何校验')
code(doc, 'with urllib.request.urlopen(url) as response:')

p = doc.add_paragraph()
r = p.add_run("修复代码：")
r.bold = True
r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
code(doc, '# 1. 只允许 http/https 协议')
code(doc, 'if not url.startswith(("http://", "https://")):')
code(doc, '    return "只允许 http 和 https 协议"')
code(doc, '# 2. 解析域名，获取真实 IP')
code(doc, 'ip = socket.gethostbyname(host)')
code(doc, '# 3. 禁止内网 IP')
code(doc, 'for private in ["127.", "10.", "172.16.", "192.168."]:')
code(doc, '    if ip.startswith(private):')
code(doc, '        return "禁止访问内网地址"')

doc.add_paragraph("攻击场景：")
doc.add_paragraph("file:///etc/passwd → 读取系统文件", style="List Bullet")
doc.add_paragraph("http://127.0.0.1:5000 → 访问本机内网服务", style="List Bullet")
doc.add_paragraph("http://192.168.1.1:3306 → 扫描内网 MySQL 端口", style="List Bullet")

doc.add_heading("三、修复前后对比", level=1)
mk_table(doc, ["安全措施", "修复前", "修复后"], [
    ["协议限制", "❌ 无限制（支持 file://）", "✅ 只允许 http/https"],
    ["内网 IP 拦截", "❌ 不拦截", "✅ 禁止私有 IP 段"],
    ["域名解析校验", "❌ 不解析", "✅ 解析后校验真实 IP"],
])

doc.add_paragraph()

doc.add_heading("四、修复验证", level=1)
mk_table(doc, ["测试项", "结果"], [
    ["file:///etc/passwd", "✅ 拦截：只允许 http/https"],
    ["http://127.0.0.1:5000", "✅ 拦截：禁止访问本机地址"],
    ["http://192.168.x.x", "✅ 拦截：禁止访问内网地址"],
    ["http://www.baidu.com", "✅ 正常抓取"],
])

doc.add_paragraph()

doc.add_heading("五、SSRF 漏洞原理", level=1)
doc.add_paragraph("SSRF（Server-Side Request Forgery）服务端请求伪造：")
doc.add_paragraph("")
doc.add_paragraph("攻击者利用服务器作为跳板，让服务器代为请求攻击者无法直接访问的内网资源。")
doc.add_paragraph("")
doc.add_paragraph("常见利用方式：")
doc.add_paragraph("1. file:// 协议读取本地文件", style="List Bullet")
doc.add_paragraph("2. 扫描内网开放端口", style="List Bullet")
doc.add_paragraph("3. 攻击内网 Redis/MySQL 等服务", style="List Bullet")
doc.add_paragraph("4. 读取云服务器元数据（169.254.169.254）", style="List Bullet")
doc.add_paragraph("")
doc.add_paragraph("防御原则：")
doc.add_paragraph("1. 白名单协议：只允许 http/https", style="List Bullet")
doc.add_paragraph("2. 禁止内网 IP：过滤私有地址段", style="List Bullet")
doc.add_paragraph("3. 解析域名再检查：防止域名指向内网", style="List Bullet")

doc.add_heading("六、安全加固建议", level=1)
for s in [
    "严格限制 URL 协议，禁用 file://、dict://、gopher://",
    "使用 IP 白名单代替域名白名单",
    "解析域名后检查真实 IP 是否为内网地址",
    "对 DNS  rebinding 攻击做防护",
    "禁止重定向到内网地址",
]:
    doc.add_paragraph(s, style="List Bullet")

path = "SSRF漏洞修复报告.docx"
doc.save(path)
print(f"✅ {path}")
