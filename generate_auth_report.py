"""生成越权漏洞修复报告和修改日志 Word 文档"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex
    })
    shading.append(shading_elem)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def generate_report():
    doc = Document()
    title = doc.add_heading("越权漏洞修复报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("项目：Flask 用户信息管理平台\n").font.size = Pt(11)
    p.add_run("修复日期：2026-07-09\n").font.size = Pt(11)
    run = p.add_run("修复状态：已全部修复")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    doc.add_paragraph()

    doc.add_heading("一、漏洞概述", level=1)
    doc.add_paragraph("本系统存在三处越权漏洞，攻击者可以通过修改 URL 参数或表单数据，访问其他用户的资料或操作他人账户。")
    doc.add_paragraph("漏洞根因：服务端未对用户身份和操作权限做校验，完全信任客户端传入的 user_id 参数。")

    doc.add_heading("二、漏洞详情", level=1)

    # 漏洞1
    doc.add_heading("漏洞 1：水平越权 — 任意查看他人资料", level=2)
    doc.add_paragraph("风险等级：🔴 严重", style="List Bullet")
    doc.add_paragraph("位置：/profile 路由 — user_id 参数", style="List Bullet")
    p = doc.add_paragraph()
    run = p.add_run("问题代码（修复前）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, "user_id = request.args.get(\"user_id\", \"\")")
    add_code_block(doc, "# 直接用传入的 user_id 查询数据库，未验证所有权")
    add_code_block(doc, "c.execute(\"SELECT ... FROM users WHERE id = ?\", (user_id,))")

    p = doc.add_paragraph()
    run = p.add_run("攻击场景：")
    run.bold = True
    doc.add_paragraph("1. alice 登录后访问 /profile?user_id=1", style="List Bullet")
    doc.add_paragraph("2. 服务器直接返回 admin 的资料（邮箱、手机、余额）", style="List Bullet")
    doc.add_paragraph("3. alice 无需密码即可查看 admin 的敏感信息", style="List Bullet")

    # 漏洞2
    doc.add_heading("漏洞 2：越权充值 — 操作他人账户", level=2)
    doc.add_paragraph("风险等级：🔴 严重", style="List Bullet")
    doc.add_paragraph("位置：/recharge 路由 — user_id 参数", style="List Bullet")
    p = doc.add_paragraph()
    run = p.add_run("问题代码（修复前）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, "user_id = request.form.get(\"user_id\", \"\")")
    add_code_block(doc, "# 直接用表单传入的 user_id 充值，未验证所有权")
    add_code_block(doc, "USERS[username][\"balance\"] += amount_val")

    p = doc.add_paragraph()
    run = p.add_run("攻击场景：")
    run.bold = True
    doc.add_paragraph("1. alice 登录后在充值表单中修改 user_id=1", style="List Bullet")
    doc.add_paragraph("2. 服务器给 admin 的账户增加了金额", style="List Bullet")
    doc.add_paragraph("3. alice 可以操控任意用户的余额", style="List Bullet")

    # 漏洞3
    doc.add_heading("漏洞 3：金额负值漏洞", level=2)
    doc.add_paragraph("风险等级：🟠 高危", style="List Bullet")
    doc.add_paragraph("位置：/recharge 路由 — amount 参数", style="List Bullet")
    p = doc.add_paragraph()
    run = p.add_run("问题代码（修复前）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, "amount = request.form.get(\"amount\", \"0\")")
    add_code_block(doc, "# 未检查 amount 正负，直接累加")
    add_code_block(doc, "USERS[username][\"balance\"] += float(amount)")

    p = doc.add_paragraph()
    run = p.add_run("攻击场景：")
    run.bold = True
    doc.add_paragraph("1. 用户输入 amount=-99999", style="List Bullet")
    doc.add_paragraph("2. 余额直接减少 99999，可以变成负数", style="List Bullet")

    doc.add_heading("三、修复方案", level=1)

    table = doc.add_table(rows=4, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["漏洞", "修复前", "修复后"]
    data = [
        ["水平越权", "不验证 user_id 所有权", "校验 user_id 是否属于当前用户"],
        ["越权充值", "不验证充值归属", "只能给自己的账号充值"],
        ["金额负值", "不检查 amount 正负", "amount <= 0 时拒绝充值"],
    ]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
    for i, row_data in enumerate(data, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading("四、防御越权漏洞的原则", level=1)
    doc.add_paragraph("1. 服务端校验：每次请求都验证用户是否有权限操作该资源", style="List Bullet")
    doc.add_paragraph("2. 使用 Session 数据：从 session 获取用户身份，不从 URL 参数取", style="List Bullet")
    doc.add_paragraph("3. 最小权限原则：普通用户和管理员的接口严格分离", style="List Bullet")
    doc.add_paragraph("4. 输入校验：金额、数量等参数必须校验范围", style="List Bullet")

    path = "越权漏洞修复报告.docx"
    doc.save(path)
    print(f"✅ 已生成：{path}  ({os.path.getsize(path)} bytes)")


def generate_changelog():
    doc = Document()
    title = doc.add_heading("修改日志 (Changelog)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("项目：Flask 用户信息管理平台\n").font.size = Pt(11)
    p.add_run("版本：v1.4（越权漏洞修复版）\n").font.size = Pt(11)
    p.add_run("日期：2026-07-09").font.size = Pt(11)
    doc.add_paragraph()

    doc.add_heading("一、修改内容", level=1)

    table = doc.add_table(rows=4, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["文件", "变更", "说明"]
    data = [
        ["app.py — /profile", "🔧 修复", "校验 user_id 所有权，普通用户只能看自己"],
        ["app.py — /recharge", "🔧 修复", "只能给自己的账号充值"],
        ["app.py — /recharge", "🔧 修复", "检查 amount > 0，拒绝负值充值"],
    ]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
    for i, row_data in enumerate(data, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text
    doc.add_paragraph()

    doc.add_heading("二、修复代码对比", level=1)

    doc.add_heading("变更 1：Profile 路由增加权限校验", level=2)
    p = doc.add_paragraph()
    run = p.add_run("修改前：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, "user_id = request.args.get(\"user_id\", \"\")")
    add_code_block(doc, "c.execute(\"SELECT ... FROM users WHERE id = ?\", (user_id,))")
    p = doc.add_paragraph()
    run = p.add_run("修改后：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code_block(doc, "# 先获取当前登录用户的 ID")
    add_code_block(doc, "c.execute(\"SELECT id FROM users WHERE username = ?\", (session[\"username\"],))")
    add_code_block(doc, "current_user_id = str(row[0])")
    add_code_block(doc, "# 校验：不是管理员且不是自己 → 拒绝")
    add_code_block(doc, "if user_id != current_user_id and not is_admin:")
    add_code_block(doc, "    return render_template(\"profile.html\", error=\"无权查看\")")

    doc.add_heading("变更 2：Recharge 路由增加权限和金额校验", level=2)
    p = doc.add_paragraph()
    run = p.add_run("修改前：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, "user_id = request.form.get(\"user_id\", \"\")")
    add_code_block(doc, "amount = request.form.get(\"amount\", \"0\")")
    add_code_block(doc, "USERS[username][\"balance\"] += float(amount)")
    p = doc.add_paragraph()
    run = p.add_run("修改后：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code_block(doc, "# 只能给自己充值")
    add_code_block(doc, "if user_id != current_user_id:")
    add_code_block(doc, "    return redirect(...)")
    add_code_block(doc, "# 检查金额必须为正")
    add_code_block(doc, "if amount_val <= 0:")
    add_code_block(doc, "    return redirect(...)")

    doc.add_heading("三、修复验证", level=1)
    tests = [
        ("普通用户查看自己", "✅ 成功"),
        ("普通用户越权看管理员", "✅ 拦截：无权查看"),
        ("给自己充负值", "✅ 拦截：金额无效"),
        ("给他人充值", "✅ 拦截：无权操作"),
        ("正常充值", "✅ 成功增加余额"),
        ("管理员查看所有用户", "✅ 允许（管理特权）"),
    ]
    ttable = doc.add_table(rows=len(tests)+1, cols=2)
    ttable.style = "Light Grid Accent 1"
    for j, h in enumerate(["测试项", "结果"]):
        cell = ttable.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
    for i, (test, result) in enumerate(tests, 1):
        ttable.rows[i].cells[0].text = test
        ttable.rows[i].cells[1].text = result
        if "✅" in result:
            set_cell_shading(ttable.rows[i].cells[1], "E8F8F5")

    doc.add_paragraph()

    doc.add_heading("四、版本历史", level=1)
    vtable = doc.add_table(rows=6, cols=4)
    vtable.style = "Light Grid Accent 1"
    vheaders = ["版本", "日期", "变更内容", "状态"]
    vdata = [
        ["v1.0", "07-07", "初始版本", "❌ 12个漏洞"],
        ["v1.1", "07-07", "密码/CSRF/Session加固", "⚠️ 仍有SQL注入"],
        ["v1.2", "07-08", "修复SQL注入", "⚠️ 文件上传漏洞"],
        ["v1.3", "07-09", "修复文件上传漏洞", "⚠️ 越权漏洞"],
        ["v1.4", "07-09", "修复越权漏洞 ✅", "✅ 当前版本"],
    ]
    for j, h in enumerate(vheaders):
        cell = vtable.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
    for i, row_data in enumerate(vdata, 1):
        for j, text in enumerate(row_data):
            vtable.rows[i].cells[j].text = text
            if "✅" in text:
                set_cell_shading(vtable.rows[i].cells[j], "E8F8F5")
            elif "❌" in text:
                set_cell_shading(vtable.rows[i].cells[j], "FDEDEC")

    path = "修改日志_CHANGELOG.docx"
    doc.save(path)
    print(f"✅ 已生成：{path}  ({os.path.getsize(path)} bytes)")


if __name__ == "__main__":
    print("=" * 50)
    print("  正在生成 Word 文档...")
    print("=" * 50)
    generate_report()
    generate_changelog()
    print()
    print("=" * 50)
    print("  全部生成完成！")
    print("=" * 50)
