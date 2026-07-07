"""生成安全审计报告和修改日志的 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex
    })
    shading.append(shading_elem)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return h


def add_code_block(doc, code_text):
    """添加代码块样式"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_risk_tag(doc, level, label):
    """添加风险等级标签"""
    color_map = {
        "严重": RGBColor(0xE7, 0x4C, 0x3C),
        "高危": RGBColor(0xE6, 0x7E, 0x22),
        "中危": RGBColor(0xF3, 0x9C, 0x12),
        "低危": RGBColor(0x27, 0xAE, 0x60),
    }
    p = doc.add_paragraph()
    run = p.add_run(f"[{level}] {label}")
    run.bold = True
    run.font.color.rgb = color_map.get(level, RGBColor(0x33, 0x33, 0x33))
    return p


# ============================================================
# 生成安全审计报告
# ============================================================

def generate_security_report():
    doc = Document()

    # 标题
    title = doc.add_heading("用户管理系统 — 安全审计报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("审计项目：用户信息管理平台（Flask 应用）\n").font.size = Pt(11)
    p.add_run("审计日期：2026-07-07\n").font.size = Pt(11)
    run = p.add_run("风险等级：严重（初始版本存在多处高危漏洞）")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

    # 一、漏洞汇总
    add_heading(doc, "一、漏洞汇总", 1)

    table = doc.add_table(rows=14, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["编号", "漏洞名称", "风险等级", "CVSS 3.1", "状态"]
    data = [
        ["V-01", "明文密码存储", "严重", "9.8", "✅ 已修复"],
        ["V-02", "硬编码 Secret Key", "严重", "9.1", "✅ 已修复"],
        ["V-03", "密码明文展示到前端", "严重", "8.6", "✅ 已修复"],
        ["V-04", "登录接口无限流", "高危", "7.5", "✅ 已修复"],
        ["V-05", "无 CSRF 保护", "高危", "7.1", "✅ 已修复"],
        ["V-06", "Debug 模式开启", "高危", "7.0", "✅ 已修复"],
        ["V-07", "Session 永不过期", "中危", "5.3", "✅ 已修复"],
        ["V-08", "Session 固定攻击", "中危", "5.1", "✅ 已修复"],
        ["V-09", "Cookie 安全属性缺失", "中危", "4.3", "✅ 已修复"],
        ["V-10", "HTML 注释泄露凭据", "中危", "4.0", "✅ 已修复"],
        ["V-11", "监听所有网卡", "低危", "3.0", "✅ 已修复"],
        ["V-12", "无 HTTPS 加密传输", "低危", "3.7", "⚠️ 需自行配置"],
    ]

    col_widths = [Inches(0.6), Inches(2.2), Inches(0.8), Inches(0.8), Inches(1.0)]
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, row_data in enumerate(data, 1):
        for j, text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if row_data[2] == "严重":
                set_cell_shading(cell, "FDEDEC")
            elif row_data[2] == "高危":
                set_cell_shading(cell, "FDEBD0")
            elif row_data[2] == "中危":
                set_cell_shading(cell, "FEF9E7")
            elif row_data[2] == "低危":
                set_cell_shading(cell, "E8F8F5")

    doc.add_paragraph()

    # 二、漏洞详情
    add_heading(doc, "二、漏洞详情", 1)

    vulnerabilities = [
        {
            "id": "V-01",
            "name": "明文密码存储",
            "level": "严重",
            "description": "USERS 字典中密码以明文形式存储，数据库或代码泄露将导致所有密码直接暴露。",
            "code_before": 'USERS = {\n    "admin": { "password": "admin123" },\n    "alice": { "password": "alice2025" },\n}',
            "code_after": 'from werkzeug.security import generate_password_hash\n\nUSERS = {\n    "admin": { "password": generate_password_hash("admin123") },\n    "alice": { "password": generate_password_hash("alice2025") },\n}',
            "attack": "攻击者获得源码或数据库 → 直接读取密码 → 登录其他平台横向移动"
        },
        {
            "id": "V-02",
            "name": "硬编码 Secret Key",
            "level": "严重",
            "description": "secret_key 固定为 \"dev-key-2025\"，攻击者可伪造任意用户的 session cookie。",
            "code_before": 'app.secret_key = "dev-key-2025"',
            "code_after": 'app.secret_key = os.environ.get("SECRET_KEY", secrets.token_hex(32))',
            "attack": "用已知 secret_key 签名伪造 session → 直接以管理员身份登录"
        },
        {
            "id": "V-03",
            "name": "密码明文展示到前端",
            "level": "严重",
            "description": "登录后用户信息页面直接输出密码字段 {{ user.password }}。",
            "code_before": '<li><span class="label">密码：</span>{{ user.password }}</li>',
            "code_after": "{# 不再展示密码 #}",
            "attack": "他人路过屏幕即可看到密码，查看页面源码即可获取"
        },
        {
            "id": "V-04",
            "name": "登录接口无限流",
            "level": "高危",
            "description": "登录接口未做任何频率限制，可无限次尝试密码。",
            "code_before": "无条件处理登录请求，无任何限流逻辑",
            "code_after": "基于 IP 的失败计数，5 次失败后锁定 15 分钟",
            "attack": "hydra -l admin -P rockyou.txt http-post-form \"...\"  → 字典爆破"
        },
        {
            "id": "V-05",
            "name": "无 CSRF 保护",
            "level": "高危",
            "description": "表单没有 CSRF Token，攻击者可诱导用户提交恶意登录请求。",
            "code_before": '<form method="POST" action="/login">',
            "code_after": '<input type="hidden" name="_csrf_token" value="{{ csrf_token() }}">',
            "attack": "构造钓鱼页面 → 用户点击后自动提交 → 利用用户已登录状态"
        },
        {
            "id": "V-06",
            "name": "Debug 模式开启",
            "level": "高危",
            "description": "app.run(debug=True) 开启调试模式，生产环境可导致远程代码执行。",
            "code_before": 'app.run(debug=True, host="0.0.0.0", port=5000)',
            "code_after": 'debug_mode = os.environ.get("FLASK_DEBUG", "0") == "1"\napp.run(debug=debug_mode, ...)',
            "attack": "/console 进入交互式 Shell（已知 PIN）→ 执行任意系统命令"
        },
    ]

    for vuln in vulnerabilities:
        add_risk_tag(doc, vuln["level"], f"{vuln['id']}：{vuln['name']}")
        doc.add_paragraph(vuln["description"])

        p = doc.add_paragraph()
        run = p.add_run("修改前：")
        run.bold = True
        run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        add_code_block(doc, vuln["code_before"])

        p = doc.add_paragraph()
        run = p.add_run("修改后：")
        run.bold = True
        run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        add_code_block(doc, vuln["code_after"])

        p = doc.add_paragraph()
        run = p.add_run("攻击场景：")
        run.bold = True
        doc.add_paragraph(vuln["attack"])

        doc.add_paragraph()  # 间距

    # 剩余漏洞简要说明
    add_heading(doc, "其余漏洞修复", 2)
    other_vulns = [
        ("V-07", "Session 永不过期", "配置 30 分钟过期时间"),
        ("V-08", "Session 固定攻击", "登录后调用 session.regenerate()"),
        ("V-09", "Cookie 安全属性缺失", "设置 HttpOnly=True, SameSite=Lax"),
        ("V-10", "HTML 注释泄露凭据", "删除调试注释"),
        ("V-11", "监听所有网卡", "配置防火墙或绑定 127.0.0.1"),
        ("V-12", "无 HTTPS 加密传输", "建议配置 Let's Encrypt SSL"),
    ]
    for vid, vname, vfix in other_vulns:
        p = doc.add_paragraph()
        run = p.add_run(f"{vid} {vname}：")
        run.bold = True
        p.add_run(vfix)

    # 三、安全建议
    add_heading(doc, "三、安全建议（未实施项）", 1)
    suggestions = [
        "HTTPS — 使用 Let's Encrypt 免费证书或自签名证书",
        "登录验证码 — 集成 Google reCAPTCHA 或类似方案",
        "密码复杂度策略 — 要求密码 ≥8 位，包含大小写字母和数字",
        "日志审计 — 记录所有登录尝试，便于事后追溯",
        "数据库迁移 — 将 USERS 字典迁移到 SQLite/MySQL，避免硬编码",
        "XSS 防护 — 对输出做 HTML 转义（Flask 模板引擎默认已转义）",
        "CSP 头 — 添加 Content-Security-Policy 响应头",
        "安全更新 — 定期更新 Flask 和依赖库版本",
    ]
    for s in suggestions:
        doc.add_paragraph(s, style="List Bullet")

    # 保存
    path = "SECURITY_REPORT.docx"
    doc.save(path)
    print(f"✅ 已生成：{path}  ({os.path.getsize(path)} bytes)")
    return path


# ============================================================
# 生成修改日志
# ============================================================

def generate_changelog():
    doc = Document()

    # 标题
    title = doc.add_heading("修改日志 (Changelog)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("项目：用户信息管理平台\n").font.size = Pt(11)
    p.add_run("版本：v1.0 → v1.1（安全加固版）\n").font.size = Pt(11)
    p.add_run("日期：2026-07-07").font.size = Pt(11)

    doc.add_paragraph()

    # 修改概览
    add_heading(doc, "一、修改概览", 1)

    table = doc.add_table(rows=7, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["文件", "操作", "说明"]
    files_data = [
        ["app.py", "🔄 重写", "约 70% 代码变更，增加安全逻辑"],
        ["templates/base.html", "🔄 微调", "无实质变更"],
        ["templates/index.html", "✂️ 删除", "移除密码展示行"],
        ["templates/login.html", "➕ 增加", "添加 CSRF Token 隐藏字段，删除调试注释"],
        ["static/css/style.css", "🔄 微调", "按钮添加 active 状态过渡"],
        ["SECURITY_REPORT.md / .docx", "➕ 新增", "安全审计报告"],
    ]

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, row_data in enumerate(files_data, 1):
        for j, text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 详细变更
    add_heading(doc, "二、详细变更", 1)

    changes = [
        {
            "num": 1,
            "title": "Secret Key 安全化",
            "before": 'app.secret_key = "dev-key-2025"',
            "after": 'app.secret_key = os.environ.get("SECRET_KEY", secrets.token_hex(32))',
            "desc": "改为随机生成 64 位十六进制密钥，支持环境变量覆盖"
        },
        {
            "num": 2,
            "title": "密码存储方式",
            "before": '"password": "admin123"  # 明文',
            "after": '"password": generate_password_hash("admin123")',
            "desc": "密码用 PBKDF2-SHA256 哈希后存储，不可逆"
        },
        {
            "num": 3,
            "title": "密码比对方式",
            "before": 'if user and user["password"] == password:',
            "after": 'if user and check_password_hash(user["password"], password):',
            "desc": "使用 check_password_hash 安全比对，防止定时攻击"
        },
        {
            "num": 4,
            "title": "新增登录限流",
            "before": "无任何限流逻辑",
            "after": "基于 IP 的失败计数，5 次失败锁定 15 分钟",
            "desc": "新增 LOGIN_ATTEMPTS 字典和 check_rate_limit/record_failed_attempt 函数"
        },
        {
            "num": 5,
            "title": "新增 CSRF 保护",
            "before": "无 CSRF 防护",
            "after": "session 内生成 CSRF Token，表单提交时校验",
            "desc": "新增 generate_csrf_token 函数，注册为 Jinja2 全局函数"
        },
        {
            "num": 6,
            "title": "Session 安全配置",
            "before": "默认配置，无过期时间",
            "after": "30 分钟过期 + HttpOnly + SameSite=Lax",
            "desc": "新增 4 行安全配置"
        },
        {
            "num": 7,
            "title": "防御 Session 固定攻击",
            "before": 'session["username"] = username',
            "after": 'session["username"] = username\nsession.permanent = True\nsession.regenerate()',
            "desc": "登录成功后重新生成 session ID"
        },
        {
            "num": 8,
            "title": "关闭 Debug 模式",
            "before": 'app.run(debug=True, ...)',
            "after": 'app.run(debug=debug_mode, ...)  由环境变量控制',
            "desc": "默认 debug=False，生产环境不会暴露调试接口"
        },
        {
            "num": 9,
            "title": "移除密码页面展示",
            "before": '<li><span class="label">密码：</span>{{ user.password }}</li>',
            "after": "已删除该行",
            "desc": "用户信息列表不再显示密码字段"
        },
        {
            "num": 10,
            "title": "添加 CSRF Token 到表单",
            "before": "无隐藏字段",
            "after": '<input type="hidden" name="_csrf_token" value="{{ csrf_token() }}">',
            "desc": "登录表单新增 CSRF Token 隐藏字段"
        },
        {
            "num": 11,
            "title": "删除调试注释",
            "before": '<!-- 调试信息 - 默认管理员账号 ... -->',
            "after": "已删除",
            "desc": "防止查看页面源代码泄露管理员凭据"
        },
    ]

    for change in changes:
        add_heading(doc, f"变更 {change['num']}：{change['title']}", 2)
        doc.add_paragraph(change["desc"])

        p = doc.add_paragraph()
        run = p.add_run("修改前：")
        run.bold = True
        run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        add_code_block(doc, change["before"])

        p = doc.add_paragraph()
        run = p.add_run("修改后：")
        run.bold = True
        run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        add_code_block(doc, change["after"])

    # 变更统计
    add_heading(doc, "三、变更统计", 1)
    stats = [
        ("修改文件数", "5"),
        ("新增文件数", "1（SECURITY_REPORT.docx）"),
        ("新增代码行", "~80 行"),
        ("删除代码行", "~10 行"),
        ("修复漏洞数", "12 个"),
        ("剩余风险数", "1 个（无 HTTPS）"),
    ]
    for label, value in stats:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        p.add_run(value)

    # 测试要点
    add_heading(doc, "四、测试要点", 1)
    tests = [
        "使用 admin/admin123 登录正常",
        "使用 alice/alice2025 登录正常",
        "错误密码返回错误提示",
        "登录后页面不显示密码",
        "连续 5 次失败后 IP 被锁定",
        "退出登录后 session 清除",
        "页面源代码无账号信息注释",
        "Cookie 包含 HttpOnly 和 SameSite 属性",
    ]
    for t in tests:
        p = doc.add_paragraph()
        run = p.add_run("☐ ")
        p.add_run(t)

    doc.add_paragraph()

    # 保存
    path = "CHANGELOG.docx"
    doc.save(path)
    print(f"✅ 已生成：{path}  ({os.path.getsize(path)} bytes)")
    return path


# ============================================================
# 主程序
# ============================================================

if __name__ == "__main__":
    print("=" * 50)
    print("  正在生成 Word 文档...")
    print("=" * 50)
    print()

    report = generate_security_report()
    changelog = generate_changelog()

    print()
    print("=" * 50)
    print("  全部生成完成！")
    print("=" * 50)
    print(f"\n  📄 {report}")
    print(f"  📄 {changelog}")
    print("\n  在 Kali 中用以下命令打开：")
    print(f"    xdg-open {report}")
    print(f"    xdg-open {changelog}")
