"""生成 SQL 注入漏洞报告和修改日志的 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex
    })
    shading.append(shading_elem)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


# ============================================================
# 漏洞报告
# ============================================================

def generate_vuln_report():
    doc = Document()

    title = doc.add_heading("SQL 注入漏洞修复报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("项目：Flask 用户信息管理平台\n").font.size = Pt(11)
    p.add_run("修复日期：2026-07-08\n").font.size = Pt(11)
    run = p.add_run("修复状态：已全部修复")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    doc.add_paragraph()

    # 一、漏洞概述
    doc.add_heading("一、漏洞概述", level=1)
    doc.add_paragraph("本系统之前存在两处 SQL 注入漏洞，攻击者可以通过在输入框中构造恶意 SQL 语句，"
                       "绕过身份验证或获取数据库中的任意数据。")
    doc.add_paragraph("漏洞根因：代码使用 f-string 将用户输入直接拼接到 SQL 语句中，"
                       "未对用户输入做任何过滤或转义，导致攻击者可以闭合 SQL 语句并注入恶意代码。")

    # 二、漏洞详情
    doc.add_heading("二、漏洞详情", level=1)

    # 漏洞1
    doc.add_heading("漏洞 1：搜索功能 SQL 注入", level=2)
    doc.add_paragraph("位置：/search 路由 — keyword 参数", style="List Bullet")
    doc.add_paragraph("风险等级：🔴 严重", style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("漏洞代码（修复前）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, 'executed_sql = f"SELECT * FROM users WHERE username LIKE \'%{keyword}%\' OR email LIKE \'%{keyword}%\'"')
    add_code_block(doc, 'c.execute(executed_sql)')

    p = doc.add_paragraph()
    run = p.add_run("修复代码（参数化查询）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code_block(doc, 'executed_sql = "SELECT * FROM users WHERE username LIKE ? OR email LIKE ?"')
    add_code_block(doc, 'c.execute(executed_sql, (f"%{keyword}%", f"%{keyword}%"))')

    doc.add_paragraph("攻击方式：")
    doc.add_paragraph("输入 ' OR '1'='1 可查询全部用户数据", style="List Bullet")
    doc.add_paragraph("输入 ' UNION SELECT 1,'inj','inj@x.com','138'-- 可插入任意伪造数据", style="List Bullet")

    # 漏洞2
    doc.add_heading("漏洞 2：注册功能 SQL 注入", level=2)
    doc.add_paragraph("位置：/register 路由 — username 参数", style="List Bullet")
    doc.add_paragraph("风险等级：🔴 严重", style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("漏洞代码（修复前）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, "sql = f\"INSERT INTO users (...) VALUES ('{username}', '{password}', '{email}', '{phone}')\"")
    add_code_block(doc, "c.execute(sql)")

    p = doc.add_paragraph()
    run = p.add_run("修复代码（参数化查询）：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code_block(doc, 'sql = "INSERT INTO users (...) VALUES (?, ?, ?, ?)"')
    add_code_block(doc, 'c.execute(sql, (username, password, email, phone))')

    doc.add_paragraph("攻击方式：")
    doc.add_paragraph("在用户名字段输入 hacker', 'pass', 'h@x.com', '123')-- 可控制写入数据库的数据", style="List Bullet")

    # 三、修复对比
    doc.add_heading("三、修复前后对比", level=1)

    table = doc.add_table(rows=3, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["对比项", "修复前（f-string 拼接）", "修复后（参数化查询）", "安全性"]
    data = [
        ["搜索功能", "f\"...LIKE '%{keyword}%'\"", "LIKE ? + 参数元组", "✅ 安全"],
        ["注册功能", "f\"...VALUES ('{username}'...)\"", "VALUES (?, ?, ?, ?)", "✅ 安全"],
    ]

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    for i, row_data in enumerate(data, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    # 四、SQL 注入原理
    doc.add_heading("四、SQL 注入原理简述", level=1)
    doc.add_paragraph("SQL 注入的本质：用户输入被当作 SQL 代码执行，而不是数据。")
    doc.add_paragraph("")
    doc.add_paragraph("漏洞条件：")
    doc.add_paragraph("1. 用户输入直接拼接到 SQL 语句中", style="List Bullet")
    doc.add_paragraph("2. 未对用户输入做过滤、转义或参数化处理", style="List Bullet")
    doc.add_paragraph("")
    doc.add_paragraph("修复原理：")
    doc.add_paragraph("参数化查询（Prepared Statement）让数据库严格区分代码和数据。"
                       "用户输入通过 ? 占位符传递，数据库会将其自动转义，不会解析其中的 SQL 关键字。")

    # 五、安全建议
    doc.add_heading("五、安全加固建议", level=1)
    suggestions = [
        "始终使用参数化查询，杜绝字符串拼接 SQL",
        "使用 ORM 框架（如 SQLAlchemy）自动处理参数化",
        "定期进行代码安全审查，扫描 SQL 注入漏洞",
        "部署 WAF（Web 应用防火墙）作为额外防护层",
        "遵循最小权限原则，数据库账号只授予必要权限",
    ]
    for s in suggestions:
        doc.add_paragraph(s, style="List Bullet")

    path = "SQL注入漏洞修复报告.docx"
    doc.save(path)
    print(f"✅ 已生成：{path}  ({os.path.getsize(path)} bytes)")
    return path


# ============================================================
# 修改日志
# ============================================================

def generate_changelog():
    doc = Document()

    title = doc.add_heading("修改日志 (Changelog)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("项目：Flask 用户信息管理平台\n").font.size = Pt(11)
    p.add_run("版本：v1.2（SQL 注入修复版）\n").font.size = Pt(11)
    p.add_run("日期：2026-07-08").font.size = Pt(11)

    doc.add_paragraph()

    # 一、修改概览
    doc.add_heading("一、修改概览", level=1)

    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["文件", "操作", "说明"]
    files_data = [
        ["app.py — /search 路由", "🔧 修复", "参数化查询替代 f-string 拼接"],
        ["app.py — /register 路由", "🔧 修复", "参数化查询替代 f-string 拼接"],
    ]

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    for i, row_data in enumerate(files_data, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    # 二、详细变更
    doc.add_heading("二、详细变更", level=1)

    # 变更1
    doc.add_heading("变更 1：搜索功能 — 参数化查询", level=2)
    doc.add_paragraph("修复位置：app.py 第 232-239 行 /search 路由")

    p = doc.add_paragraph()
    run = p.add_run("修改前：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, "executed_sql = f\"SELECT ... WHERE username LIKE '%{keyword}%' OR email LIKE '%{keyword}%'\"")
    add_code_block(doc, "c.execute(executed_sql)")

    p = doc.add_paragraph()
    run = p.add_run("修改后：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code_block(doc, 'executed_sql = "SELECT ... WHERE username LIKE ? OR email LIKE ?"')
    add_code_block(doc, 'c.execute(executed_sql, (f"%{keyword}%", f"%{keyword}%"))')

    doc.add_paragraph("说明：通过 ? 占位符传递用户输入，数据库自动转义特殊字符，杜绝 SQL 注入。")

    # 变更2
    doc.add_heading("变更 2：注册功能 — 参数化查询", level=2)
    doc.add_paragraph("修复位置：app.py 第 197-203 行 /register 路由")

    p = doc.add_paragraph()
    run = p.add_run("修改前：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    add_code_block(doc, 'sql = f"INSERT INTO users (...) VALUES (\'{username}\', \'{password}\', \'{email}\', \'{phone}\')"')
    add_code_block(doc, "c.execute(sql)")

    p = doc.add_paragraph()
    run = p.add_run("修改后：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    add_code_block(doc, 'sql = "INSERT INTO users (username, password, email, phone) VALUES (?, ?, ?, ?)"')
    add_code_block(doc, 'c.execute(sql, (username, generate_password_hash(password), email, phone))')

    doc.add_paragraph("说明：使用参数化查询，用户输入通过元组传递，无法闭合 SQL 语句。")

    # 三、修复验证
    doc.add_heading("三、修复验证", level=1)
    doc.add_paragraph("修复后使用以下 payload 测试，均不再生效：")

    tests = [
        "搜索 ' OR '1'='1 → 不再返回全部用户",
        "搜索 ' UNION SELECT 1,'inj','inj@x.com','138'-- → 不再插入伪造数据",
        "注册用户名 hacker', 'pass', 'h@x.com', '123')-- → 不再注入自定义数据",
        "搜索正常关键词（如 admin）→ 仍能正常查询",
        "注册正常用户 → 仍能正常注册并登录",
    ]
    for t in tests:
        doc.add_paragraph(t, style="List Bullet")

    # 四、版本记录
    doc.add_heading("四、版本历史", level=1)

    vtable = doc.add_table(rows=4, cols=4)
    vtable.style = "Light Grid Accent 1"
    vtable.alignment = WD_TABLE_ALIGNMENT.CENTER

    vheaders = ["版本", "日期", "变更内容", "状态"]
    vdata = [
        ["v1.0", "2026-07-07", "初始版本（包含12个安全漏洞）", "❌ 不安全"],
        ["v1.1", "2026-07-07", "安全加固版（修复密码、CSRF等）", "⚠️ 仍有SQL注入"],
        ["v1.2", "2026-07-08", "修复SQL注入漏洞（当前版本）", "✅ 安全"],
    ]

    for j, h in enumerate(vheaders):
        cell = vtable.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2C3E50")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    for i, row_data in enumerate(vdata, 1):
        for j, text in enumerate(row_data):
            vtable.rows[i].cells[j].text = text

    path = "修改日志_CHANGELOG.docx"
    doc.save(path)
    print(f"✅ 已生成：{path}  ({os.path.getsize(path)} bytes)")
    return path


if __name__ == "__main__":
    print("=" * 50)
    print("  正在生成 Word 文档...")
    print("=" * 50)
    report = generate_vuln_report()
    changelog = generate_changelog()
    print()
    print("=" * 50)
    print("  全部生成完成！")
    print("=" * 50)
    print(f"\n  📄 {report}")
    print(f"  📄 {changelog}")
