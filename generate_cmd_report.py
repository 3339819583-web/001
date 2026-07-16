"""生成命令注入漏洞修复报告"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def sc(cell, color):
    s = cell._element.get_or_add_tcPr()
    e = s.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color})
    s.append(e)


def code(doc, text, color=RGBColor(0x8E, 0x44, 0xAD)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = color


def mk_table(doc, headers, data):
    t = doc.add_table(rows=len(data)+1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        sc(c, "2C3E50")
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
            if "✅" in val:
                sc(t.rows[i].cells[j], "E8F8F5")
            elif "❌" in val:
                sc(t.rows[i].cells[j], "FDEDEC")
    return t


doc = Document()
t = doc.add_heading("命令注入漏洞修复报告", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("项目：Flask 用户信息管理平台\n").font.size = Pt(11)
p.add_run("修复日期：2026-07-14\n").font.size = Pt(11)
r = p.add_run("修复状态：已全部修复")
r.bold = True
r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
doc.add_paragraph()

doc.add_heading("一、漏洞概述", level=1)
doc.add_paragraph("Ping 网络诊断功能存在命令注入漏洞，攻击者可以通过在 IP 地址中插入特殊字符（如 |、;、&& 等）执行任意系统命令，获取服务器的完全控制权。")

doc.add_heading("二、漏洞详情", level=1)
doc.add_heading("漏洞：命令注入", level=2)
doc.add_paragraph("风险等级：🔴 严重", style="List Bullet")
doc.add_paragraph("位置：/ping 路由 — ip 参数", style="List Bullet")

p = doc.add_paragraph()
r = p.add_run("问题代码（修复前）：")
r.bold = True
r.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
code(doc, 'command = f"ping -c 3 {ip_input}"               # f-string 拼接')
code(doc, 'subprocess.check_output(command, shell=True)     # shell=True')
doc.add_paragraph("")
p = doc.add_paragraph()
r = p.add_run("修复代码：")
r.bold = True
r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
code(doc, '# 1. 输入校验：只允许字母数字和.-')
code(doc, 'if not re.match(r\'^[a-zA-Z0-9\\.\\-]+$\', ip_input):')
code(doc, '    error = "格式不合法"')
code(doc, '# 2. 参数列表方式执行，禁用 shell=True')
code(doc, 'subprocess.check_output(["ping", "-c", "3", ip_input])')

doc.add_paragraph("攻击场景：")
doc.add_paragraph('127.0.0.1 | hostname → 执行 hostname 命令查看主机名', style="List Bullet")
doc.add_paragraph('127.0.0.1 ; id → 查看当前用户 UID 信息', style="List Bullet")
doc.add_paragraph('127.0.0.1 && cat /etc/passwd → 读取系统文件', style="List Bullet")

doc.add_heading("三、修复前后对比", level=1)
mk_table(doc, ["安全措施", "修复前", "修复后"], [
    ["命令拼接方式", "❌ f-string 拼接", "✅ 参数列表传递"],
    ["shell", "❌ shell=True", "✅ shell=False"],
    ["输入校验", "❌ 无校验", "✅ 正则限制合法字符"],
])

doc.add_paragraph()

doc.add_heading("四、修复验证", level=1)
mk_table(doc, ["测试项", "结果"], [
    ["正常 ping 8.8.8.8", "✅ 正常执行"],
    ['命令注入：127.0.0.1;id', "✅ 拦截：格式不合法"],
    ['命令注入：127.0.0.1|whoami', "✅ 拦截：格式不合法"],
    ['命令注入：127.0.0.1 && cat /etc/passwd', "✅ 拦截：格式不合法"],
])

doc.add_paragraph()

doc.add_heading("五、命令注入漏洞原理", level=1)
doc.add_paragraph("命令注入（Command Injection）是攻击者利用系统命令拼接漏洞，在正常命令中插入额外命令的攻击方式。")
doc.add_paragraph("")
doc.add_paragraph("常见注入符号：")
doc.add_paragraph(";      → 命令分隔符，顺序执行多条命令", style="List Bullet")
doc.add_paragraph("|      → 管道符，前一条命令的输出传给下一条", style="List Bullet")
doc.add_paragraph("&&     → 前一条成功才执行后一条", style="List Bullet")
doc.add_paragraph("`cmd`  → 反引号，先执行内部命令", style="List Bullet")
doc.add_paragraph("$(cmd) → bash 语法，先执行内部命令", style="List Bullet")
doc.add_paragraph("")
doc.add_paragraph("防御原则：")
doc.add_paragraph("1. 永远不要用 shell=True", style="List Bullet")
doc.add_paragraph("2. 永远不要拼接命令字符串", style="List Bullet")
doc.add_paragraph("3. 使用参数列表传递参数", style="List Bullet")
doc.add_paragraph("4. 对用户输入做严格校验", style="List Bullet")

doc.add_heading("六、版本历史", level=1)
mk_table(doc, ["版本", "日期", "变更", "状态"], [
    ["v1.0", "07-07", "初始版本", "❌"],
    ["v1.1", "07-07", "密码/CSRF加固", "⚠️"],
    ["v1.2", "07-08", "修复SQL注入", "⚠️"],
    ["v1.3", "07-09", "修复文件上传", "⚠️"],
    ["v1.4", "07-09", "修复越权", "⚠️"],
    ["v1.5", "07-13", "修复LFI", "⚠️"],
    ["v1.6", "07-13", "修复密码修改", "⚠️"],
    ["v1.7", "07-14", "修复SSRF", "⚠️"],
    ["v1.8", "07-14", "修复命令注入", "✅"],
])

path = "命令注入漏洞修复报告.docx"
doc.save(path)
print(f"✅ {path}")
